# -*- coding: utf-8 -*-
from docx import Document
import os
import sys

# Set UTF-8 encoding
sys.stdout.reconfigure(encoding='utf-8')

# Read the Word document
doc_path = r'D:\工作\WPS同步\(A材料模板)\AI项目\签证平台\模块化后台管理系统框架.docx'
output_path = r'C:\Users\Administrator\Desktop\muhai001\muhai001\extracted_content.txt'

# Check if file exists
if os.path.exists(doc_path):
    doc = Document(doc_path)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + '\n')
                f.write('---\n')
        
        # Extract from tables
        for table_idx, table in enumerate(doc.tables):
            f.write(f"\n=== Table {table_idx + 1} ===\n")
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    f.write(' | '.join(row_text) + '\n')
        
        f.write(f'\n=== Total paragraphs: {len(doc.paragraphs)} ===\n')
        f.write(f'=== Total tables: {len(doc.tables)} ===\n')
    
    print(f"Content extracted to: {output_path}")
else:
    print('File not found:', doc_path)
